
import os
import time
import uuid
from datetime import datetime

from flask import Flask, request, jsonify, render_template, send_from_directory
from google import genai
from docx import Document
from pypdf import PdfReader
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
GENERATED_DIR = os.path.join("/tmp", "qraft_generated")
os.makedirs(GENERATED_DIR, exist_ok=True)

client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
MODEL = "gemini-flash-latest"

K_LEVEL_DEFS = {
    "K1": "Remember - recall facts, terms, basic concepts (define, list, state, name)",
    "K2": "Understand - explain ideas or concepts (explain, describe, summarize, classify)",
    "K3": "Apply - use information in new situations (solve, demonstrate, apply, calculate)",
    "K4": "Analyze - draw connections among ideas (compare, differentiate, examine, categorize)",
    "K5": "Evaluate - justify a stand or decision (justify, critique, assess, argue)",
    "K6": "Create - produce new or original work (design, develop, formulate, propose)",
}

# Exam-type defaults shown on the landing page (Anna University style)
EXAM_DEFAULTS = {
    "CIA 1":         {"start": "K1", "end": "K3", "total_questions": 10, "marks": 5},
    "CIA 2":         {"start": "K1", "end": "K4", "total_questions": 10, "marks": 5},
    "Model Exam":    {"start": "K1", "end": "K5", "total_questions": 20, "marks": 5},
    "Semester Exam": {"start": "K1", "end": "K6", "total_questions": 20, "marks": 5},
}


# ---------------------------------------------------------------------------
# Range -> per-level distribution
# ---------------------------------------------------------------------------
def distribute_questions(start_level, end_level, total_questions):
    """Spread total_questions evenly across the K-levels from start_level to
    end_level (inclusive). Leftover questions go to the earlier levels."""
    levels_order = list(K_LEVEL_DEFS.keys())
    start_idx = levels_order.index(start_level)
    end_idx = levels_order.index(end_level)
    selected_levels = levels_order[start_idx:end_idx + 1]

    num_levels = len(selected_levels)
    base = total_questions // num_levels
    remainder = total_questions % num_levels

    k_distribution = {level: 0 for level in levels_order}
    for i, level in enumerate(selected_levels):
        k_distribution[level] = base + (1 if i < remainder else 0)

    return k_distribution
def build_section_distribution(start_level, end_level, total_questions):
    
    levels_order = list(K_LEVEL_DEFS.keys())

    start_idx = levels_order.index(start_level)
    end_idx = levels_order.index(end_level)

    selected_levels = levels_order[
        start_idx:end_idx + 1
    ]

    number_of_levels = len(selected_levels)

    base = total_questions // number_of_levels
    remainder = total_questions % number_of_levels

    distribution = {}

    for i, level in enumerate(selected_levels):

        distribution[level] = (
            base + (1 if i < remainder else 0)
        )

    return distribution

# ---------------------------------------------------------------------------
# Prompt builder
# ---------------------------------------------------------------------------
def build_prompt(core_content, subject, k_distribution, marks_per_question, exam_title):
    level_lines = []
    total_q = 0
    for level, count in k_distribution.items():
        if count and count > 0:
            level_lines.append(f"- {level} ({K_LEVEL_DEFS[level]}): {count} question(s)")
            total_q += count

    total_marks = total_q * marks_per_question

    prompt = f"""You are an experienced examination question paper setter for Indian
university/college exams. Create a formal question paper strictly based on the
CORE CONTENT provided below. Do not invent topics outside this content.

EXAM TITLE: {exam_title}
SUBJECT: {subject}

CORE CONTENT (syllabus / notes to base every question on):
\"\"\"
{core_content}
\"\"\"

Generate questions according to this Bloom's Taxonomy K-level distribution:
{chr(10).join(level_lines)}

Rules:
1. Each question must be tagged with its K-level in square brackets at the end, e.g. "[K2]".
2. Each question is worth {marks_per_question} marks. Total marks = {total_marks}.
3. Number questions sequentially (1, 2, 3, ...).
4. Group questions level by level, in the order K1 -> K6, with a small heading
   before each group like "Section A - K1 (Remember)".
5. Keep language clear and exam-appropriate. Questions must be answerable
   purely from the CORE CONTENT above.
6. Output ONLY the question paper content in this exact plain-text structure
   (no markdown, no extra commentary):

Section A - K1 (Remember)
1. <question text> [K1] (<marks> marks)
2. ...

Section B - K2 (Understand)
...

(continue only for levels that have count > 0)

At the very end add a line: TOTAL MARKS: {total_marks}
"""
    return prompt, total_marks, total_q
def build_pattern_prompt(core_content, subject, pattern, exam_title):
    
    total_marks = 0
    total_questions = 0
    pattern_details = []

    for section in pattern:

        section_name = section.get("section", "Part")
        question_count = int(section.get("question_count", 0))
        marks = int(section.get("marks", 0))
        has_choice = section.get("has_choice", False)
        k_levels = section.get("k_levels", {})

        total_questions += question_count
        total_marks += question_count * marks

        k_level_text = []

        for level in ["K1", "K2", "K3", "K4", "K5", "K6"]:

            count = int(k_levels.get(level, 0))

            if count > 0:
                k_level_text.append(
                    f"{level}: {count} question(s)"
                )

        if has_choice:
            choice_text = "YES - Every question must have A OR B choice."
        else:
            choice_text = "NO"

        pattern_details.append(f"""
SECTION: {section_name}
QUESTIONS: {question_count}
MARKS: {marks}
INTERNAL CHOICE: {choice_text}

K-LEVEL:
{chr(10).join(k_level_text)}
""")

    prompt = f"""
You are an expert university question paper setter.

Create one complete question paper.

EXAM TITLE: {exam_title}
SUBJECT: {subject}

SYLLABUS:
{core_content}

QUESTION PAPER PATTERN:

{chr(10).join(pattern_details)}

RULES:

1. Follow the exact pattern.
2. Generate the exact number of questions.
3. Follow the specified marks.
4. Follow the specified K-Level distribution.
5. Add K-Level to every question.
6. If A OR B choice is enabled:

11. (A) Question [K-Level]
    OR
    (B) Alternative Question [K-Level]

Student must answer either A or B.

7. Do not repeat questions.
8. Generate questions only from the syllabus.
9. Number questions continuously.

At the end:

TOTAL QUESTIONS: {total_questions}
TOTAL MARKS: {total_marks}
"""

    return prompt, total_marks, total_questions

# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    return render_template(
        "index.html",
        k_defs=K_LEVEL_DEFS,
        exam_defaults=EXAM_DEFAULTS,
    )


@app.route("/api/upload", methods=["POST"])
def upload_syllabus():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = file.filename.lower()

    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(file)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)

        elif filename.endswith(".docx"):
            doc = Document(file)
            text = "\n".join(p.text for p in doc.paragraphs)

        elif filename.endswith(".txt"):
            text = file.read().decode("utf-8", errors="ignore")

        else:
            return jsonify({"error": "Only PDF, DOCX or TXT files are supported"}), 400

        text = text.strip()
        if not text:
            return jsonify({"error": "Could not read any text from this file"}), 400

        return jsonify({"text": text})

    except Exception as e:
        return jsonify({"error": f"Could not read this file: {str(e)}"}), 500


@app.route("/api/generate", methods=["POST"])
def generate_question_paper():
    data = request.get_json(force=True)

    core_content = (data.get("core_content") or "").strip()
    subject = (data.get("subject") or "General").strip()
    exam_title = (data.get("exam_title") or "Internal Assessment Test").strip()
    marks_per_question = int(data.get("marks_per_question") or 5)
    start_level = data.get("start_level") or "K1"
    end_level = data.get("end_level") or "K4"
    total_questions = int(data.get("total_questions") or 0)
    file_format = (data.get("format") or "docx").lower()  # "docx" or "pdf"

    if not core_content:
        return jsonify({"error": "core_content (topic/notes) is required"}), 400
    if total_questions <= 0:
        return jsonify({"error": "total_questions must be at least 1"}), 400

    levels_order = list(K_LEVEL_DEFS.keys())
    if start_level not in levels_order or end_level not in levels_order:
        return jsonify({"error": "Invalid K-level range"}), 400
    if levels_order.index(start_level) > levels_order.index(end_level):
        return jsonify({"error": "Start level must come before end level (e.g. K1 to K4)"}), 400
    if file_format not in ("docx", "pdf"):
        file_format = "docx"

    k_distribution = distribute_questions(start_level, end_level, total_questions)

    prompt, total_marks, total_q = build_prompt(
        core_content, subject, k_distribution, marks_per_question, exam_title
    )

    # NOTE: while debugging, the real error message is returned to the
    # browser so you can see exactly what went wrong without checking the
    # terminal. Once everything works, you can change the except block back
    # to a generic message if you don't want technical details shown.
    last_error = None
    paper_text = None
    attempt = 0

    while True:
        attempt += 1
        try:
            response = client.models.generate_content(
                model=MODEL,
                contents=prompt,
            )
            paper_text = (response.text or "").strip()
            last_error = None
            break
        except Exception as e:
            last_error = e
            error_str = str(e)
            if "503" in error_str or "UNAVAILABLE" in error_str or "overloaded" in error_str.lower():
                wait_time = min(2 * attempt, 15)
                time.sleep(wait_time)
                continue
            break

    if last_error is not None:
        return jsonify({"error": f"Generation failed: {str(last_error)}"}), 500
    if not paper_text:
        return jsonify({"error": "The AI returned an empty response. Please try again."}), 500

    filename_base = f"question_paper_{uuid.uuid4().hex[:8]}"
    try:
        if file_format == "pdf":
            filename = f"{filename_base}.pdf"
            filepath = os.path.join(GENERATED_DIR, filename)
            save_as_pdf(paper_text, exam_title, subject, total_marks, filepath)
        else:
            filename = f"{filename_base}.docx"
            filepath = os.path.join(GENERATED_DIR, filename)
            save_as_docx(paper_text, exam_title, subject, total_marks, filepath)
    except Exception as e:
        return jsonify({"error": f"File creation failed: {str(e)}"}), 500

    return jsonify({
        "paper_text": paper_text,
        "total_marks": total_marks,
        "total_questions": total_q,
        "k_distribution": k_distribution,
        "download_url": f"/download/{filename}",
    })
@app.route("/api/generate-pattern", methods=["POST"])
def generate_pattern_question_paper():

    data = request.get_json(force=True)

    core_content = (
        data.get("core_content") or ""
    ).strip()

    subject = (
        data.get("subject") or "General"
    ).strip()

    exam_title = (
        data.get("exam_title")
        or "Question Paper"
    ).strip()

    pattern = data.get("pattern") or []

    file_format = (
        data.get("format") or "docx"
    ).lower()


    if not core_content:

        return jsonify({
            "error":
                "Syllabus or core content is required"
        }), 400


    if not pattern:

        return jsonify({
            "error":
                "Question paper pattern is required"
        }), 400


    levels_order = list(K_LEVEL_DEFS.keys())

    final_pattern = []

    total_questions = 0
    total_marks = 0


    # ==========================================
    # PROCESS EACH PART
    # ==========================================

    for section in pattern:

        section_name = (
            section.get("section") or "PART"
        )

        question_count = int(
            section.get("question_count") or 0
        )

        marks = int(
            section.get("marks") or 0
        )

        has_choice = bool(
            section.get("has_choice", False)
        )

        start_level = (
            section.get("start_level") or "K1"
        )

        end_level = (
            section.get("end_level") or "K1"
        )


        if question_count <= 0:

            return jsonify({
                "error":
                    f"{section_name}: "
                    "Question count must be greater than 0"
            }), 400


        if marks <= 0:

            return jsonify({
                "error":
                    f"{section_name}: "
                    "Marks must be greater than 0"
            }), 400


        if (
            start_level not in levels_order
            or end_level not in levels_order
        ):

            return jsonify({
                "error":
                    f"{section_name}: Invalid K-Level"
            }), 400


        if (
            levels_order.index(start_level)
            >
            levels_order.index(end_level)
        ):

            return jsonify({
                "error":
                    f"{section_name}: "
                    "Start K-Level must come before End K-Level"
            }), 400


        # AUTO DISTRIBUTE QUESTIONS
        k_distribution = build_section_distribution(
            start_level,
            end_level,
            question_count
        )


        final_pattern.append({

            "section":
                section_name,

            "question_count":
                question_count,

            "marks":
                marks,

            "has_choice":
                has_choice,

            "start_level":
                start_level,

            "end_level":
                end_level,

            "k_distribution":
                k_distribution

        })


        total_questions += question_count

        total_marks += (
            question_count * marks
        )


    # ==========================================
    # BUILD PROMPT
    # ==========================================

    pattern_text = []


    for section in final_pattern:

        section_lines = []

        for level in levels_order:

            count = (
                section["k_distribution"]
                .get(level, 0)
            )

            if count > 0:

                section_lines.append(
                    f"{level}: {count} question(s)"
                )


        choice_text = (
            "YES - Every numbered question "
            "must contain an A OR B alternative."
            if section["has_choice"]
            else "NO"
        )


        pattern_text.append(
            f"""
SECTION: {section["section"]}

NUMBER OF QUESTIONS:
{section["question_count"]}

MARKS PER QUESTION:
{section["marks"]}

K-LEVEL RANGE:
{section["start_level"]} TO {section["end_level"]}

AUTO DISTRIBUTION:
{chr(10).join(section_lines)}

INTERNAL CHOICE:
{choice_text}
"""
        )


    prompt = f"""
You are an expert university examination
question paper setter.

Create ONE COMPLETE QUESTION PAPER strictly
from the syllabus provided below.

EXAM TITLE:
{exam_title}

SUBJECT:
{subject}


SYLLABUS / CORE CONTENT:

\"\"\"
{core_content}
\"\"\"


QUESTION PAPER PATTERN:

{chr(10).join(pattern_text)}


STRICT RULES:

1. Follow the exact PART structure.

2. Generate the exact number of numbered
questions specified for each PART.

3. Follow the exact marks for every PART.

4. Follow the K-Level distribution exactly.

5. Every question must end with its K-Level,
for example:

Explain Python variables. [K2]

6. Do not generate questions outside the
provided syllabus.

7. Do not repeat questions.

8. Number questions continuously throughout
the paper.

9. If INTERNAL CHOICE is YES, format every
question like this:

1. (A) First question [K-Level]

   OR

   (B) Alternative question [K-Level]

The student must answer either A or B.

IMPORTANT:
A and B must test similar marks and similar
difficulty.

10. If INTERNAL CHOICE is NO, generate one
question only.

11. Use clear university examination language.

12. Keep each PART clearly separated.

13. Use this output format:


PART A

1. Question text [K-Level] (2 Marks)

2. Question text [K-Level] (2 Marks)


PART B

11. (A) Question text [K-Level] (5 Marks)

    OR

    (B) Alternative question [K-Level]
        (5 Marks)


PART C

21. Question text [K-Level] (10 Marks)


At the end write:

TOTAL QUESTIONS: {total_questions}

TOTAL MARKS: {total_marks}


Output ONLY the question paper.
Do not add explanations or commentary.
"""


    # ==========================================
    # CALL GEMINI
    # ==========================================

    last_error = None
    paper_text = None
    attempt = 0


    while True:

        attempt += 1

        try:

            response = (
                client.models.generate_content(
                    model=MODEL,
                    contents=prompt
                )
            )


            paper_text = (
                response.text or ""
            ).strip()


            last_error = None

            break


        except Exception as e:

            last_error = e

            error_str = str(e).lower()


            if (
                "503" in str(e)
                or "unavailable" in error_str
                or "overloaded" in error_str
            ):

                wait_time = min(
                    2 * attempt,
                    15
                )

                time.sleep(wait_time)

                continue


            break


    if last_error is not None:

        return jsonify({
            "error":
                f"Generation failed: {str(last_error)}"
        }), 500


    if not paper_text:

        return jsonify({
            "error":
                "AI returned an empty response"
        }), 500


    # ==========================================
    # CREATE FILE
    # ==========================================

    filename_base = (
        f"question_paper_"
        f"{uuid.uuid4().hex[:8]}"
    )


    try:

        if file_format == "pdf":

            filename = (
                f"{filename_base}.pdf"
            )

            filepath = os.path.join(
                GENERATED_DIR,
                filename
            )


            save_as_pdf(
                paper_text,
                exam_title,
                subject,
                total_marks,
                filepath
            )


        else:

            filename = (
                f"{filename_base}.docx"
            )

            filepath = os.path.join(
                GENERATED_DIR,
                filename
            )


            save_as_docx(
                paper_text,
                exam_title,
                subject,
                total_marks,
                filepath
            )


    except Exception as e:

        return jsonify({
            "error":
                f"File creation failed: {str(e)}"
        }), 500


    return jsonify({

        "paper_text":
            paper_text,

        "total_marks":
            total_marks,

        "total_questions":
            total_questions,

        "download_url":
            f"/download/{filename}"

    })
@app.route("/api/generate-pattern", methods=["POST"])
def generate_pattern_paper():

    data = request.get_json(force=True)

    core_content = (data.get("core_content") or "").strip()
    subject = (data.get("subject") or "General").strip()
    exam_title = (data.get("exam_title") or "Custom Question Paper").strip()

    pattern = data.get("pattern") or []
    file_format = (data.get("format") or "docx").lower()

    # Check syllabus
    if not core_content:
        return jsonify({
            "error": "Please provide syllabus/content"
        }), 400

    # Check pattern
    if not pattern:
        return jsonify({
            "error": "Please add at least one question pattern"
        }), 400

    # Validate each pattern section
    for section in pattern:

        question_count = int(
            section.get("question_count", 0)
        )

        marks = int(
            section.get("marks", 0)
        )

        if question_count <= 0:
            return jsonify({
                "error": f"{section.get('section', 'Section')} must have at least 1 question"
            }), 400

        if marks <= 0:
            return jsonify({
                "error": f"{section.get('section', 'Section')} marks must be greater than 0"
            }), 400

    if file_format not in ("pdf", "docx"):
        file_format = "docx"

    # Build the custom pattern prompt
    prompt, total_marks, total_questions = build_pattern_prompt(
        core_content,
        subject,
        pattern,
        exam_title
    )

    # Generate questions using Gemini
    last_error = None
    paper_text = None
    attempt = 0

    while True:
        attempt += 1

        try:
            response = client.models.generate_content(
                model=MODEL,
                contents=prompt,
            )

            paper_text = (response.text or "").strip()
            last_error = None
            break

        except Exception as e:
            last_error = e
            error_str = str(e)

            if (
                "503" in error_str
                or "UNAVAILABLE" in error_str
                or "overloaded" in error_str.lower()
            ):
                wait_time = min(2 * attempt, 15)
                time.sleep(wait_time)
                continue

            break

    if last_error is not None:
        return jsonify({
            "error": f"Pattern generation failed: {str(last_error)}"
        }), 500

    if not paper_text:
        return jsonify({
            "error": "The AI returned an empty response. Please try again."
        }), 500

    # Create PDF or DOCX
    filename_base = f"pattern_question_paper_{uuid.uuid4().hex[:8]}"

    try:

        if file_format == "pdf":

            filename = f"{filename_base}.pdf"
            filepath = os.path.join(GENERATED_DIR, filename)

            save_as_pdf(
                paper_text,
                exam_title,
                subject,
                total_marks,
                filepath
            )

        else:

            filename = f"{filename_base}.docx"
            filepath = os.path.join(GENERATED_DIR, filename)

            save_as_docx(
                paper_text,
                exam_title,
                subject,
                total_marks,
                filepath
            )

    except Exception as e:
        return jsonify({
            "error": f"File creation failed: {str(e)}"
        }), 500

    return jsonify({
        "paper_text": paper_text,
        "total_marks": total_marks,
        "total_questions": total_questions,
        "download_url": f"/download/{filename}",
    })


@app.route("/download/<path:filename>")
def download(filename):
    return send_from_directory(GENERATED_DIR, filename, as_attachment=True)


# ---------------------------------------------------------------------------
# DOCX writer
# ---------------------------------------------------------------------------
def save_as_docx(paper_text, exam_title, subject, total_marks, filepath):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(exam_title)
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"Subject: {subject}    |    Total Marks: {total_marks}    |    Date: {datetime.now().strftime('%d-%m-%Y')}"
    )
    sub_run.font.size = Pt(11)

    doc.add_paragraph()

    for line in paper_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.lower().startswith("section"):
            h = doc.add_paragraph()
            hr = h.add_run(line)
            hr.bold = True
            hr.font.size = Pt(13)
        elif line.upper().startswith("TOTAL MARKS"):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
        else:
            doc.add_paragraph(line)

    doc.save(filepath)


# ---------------------------------------------------------------------------
# PDF writer
# ---------------------------------------------------------------------------
def save_as_pdf(paper_text, exam_title, subject, total_marks, filepath):
    doc = SimpleDocTemplate(filepath, pagesize=A4, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "TitleCenter", parent=styles["Title"], alignment=TA_CENTER, fontSize=16
    )
    sub_style = ParagraphStyle(
        "SubCenter", parent=styles["Normal"], alignment=TA_CENTER, fontSize=10.5
    )
    section_style = ParagraphStyle(
        "Section", parent=styles["Heading2"], fontSize=12.5, spaceBefore=12
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=10.5, spaceAfter=6, leading=15
    )

    story = [
        Paragraph(exam_title, title_style),
        Paragraph(
            f"Subject: {subject} &nbsp;&nbsp;|&nbsp;&nbsp; Total Marks: {total_marks} "
            f"&nbsp;&nbsp;|&nbsp;&nbsp; Date: {datetime.now().strftime('%d-%m-%Y')}",
            sub_style,
        ),
        Spacer(1, 16),
    ]

    for line in paper_text.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        # Escape characters that would break ReportLab's mini-XML markup
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if line.lower().startswith("section"):
            story.append(Paragraph(safe_line, section_style))
        elif line.upper().startswith("TOTAL MARKS"):
            story.append(Paragraph(f"<b>{safe_line}</b>", body_style))
        else:
            story.append(Paragraph(safe_line, body_style))

    doc.build(story)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
